import pandas as pd
import os
from tprint import tprint
from docx import Document
from datetime import datetime
from docx.shared import Pt
from docx.oxml.ns import qn
import openpyxl
from openpyxl.styles import PatternFill

#
# ПАРСИНГ CSV
#

def clean_name(name):
    name = name.replace(' ', '').replace('?', '').replace('!', '').replace('P', 'Р').replace('/', '').replace('\\', '').lstrip('0')
    if 'Р' in name:
        name = f"Р-{name.replace('Р-', '').lstrip('0')}"
    return name

csv_path = os.path.join(os.path.expanduser('~'), 'Documents', '111.csv')

df = pd.read_csv(csv_path)
tprint(df)

num = df.loc[df[df.columns[0]] == 'Номер заказа'].index
cards = []
print(num)
for i, index in enumerate(num):



    print(i, index)
    start = index
    if i + 1 < len(num):
        end = num[i + 1] - 2
    else:
        end = None  # до конца таблицы

    chunk = df.loc[start:end].T.bfill()
    chunk.columns = chunk.iloc[0]
    chunk = chunk.iloc[1:].reset_index(drop=True)
    # chunk = chunk.drop(columns=['Адрес', 'Дата заказа', 'Дата заказа - год', 'Последнее состояние заказа',
    #                             'Дата заказа - месяц', 'Срок исполнения', 'Состояния заказа',
    #                             'Состояние заказа', 'Дата'])
    chunk = chunk.rename(columns={'Задание на сканирование': 'Номер фонда'})

    for j, col in enumerate(chunk.columns):
        if chunk.iloc[0, j] == col:
            chunk.iloc[0, j] = chunk.iloc[1, j]

    chunk = chunk[:-1]
    chunk.columns = [f"{col}_{j}" if chunk.columns.duplicated()[j] else col
                     for j, col in enumerate(chunk.columns)]
    row = chunk.iloc[0]  # Series
    tprint(chunk)
    #### если несколько дел дублируем
    dict_ = {
        'номер каиса': row['Номер заказа'],
        'фио': row['Заказчик'].strip().replace('  ', ' '),
        'шифры': []
    }

    for base_idx, col in enumerate(chunk.columns):
        if "Номер фонда" in col:

            # находим реальные имена колонок в "блоке"
            col_fond = chunk.columns[base_idx]
            col_opis = chunk.columns[base_idx + 1]
            col_delo = chunk.columns[base_idx + 2]

            # ищем внутри следующих колонок именно "Листы" и "Разрешение в dpi"
            col_listy = next(
                (c for c in chunk.columns[base_idx:] if "Листы" in c), None)
            col_dpi = next(
                (c for c in chunk.columns[base_idx:] if "Разрешение в dpi" in c), None)

            shifr = {
                'номер фонда': row[col_fond],
                'номер описи': row[col_opis],
                'номер дела': row[col_delo],
                'листы': row[col_listy].strip() if col_listy else None,
                'разрешение': row[col_dpi] if col_dpi else None
            }
            dict_['шифры'].append(shifr)

    print(dict_)
    cards.append(dict_)

#
# ПРОВЕРКА НАЛИЧИЯ КОПИЙ, редактирование разрешения, распределение по папкам
#

os.chdir(r'D:\Заказы\_КАИСА\_Заказ дел')



def find_directory(base_path, dir_name):
    for dir in os.listdir(base_path):
        clean_dir_name = dir.replace('(', ' ').replace('з', ' ').replace('z', ' ').split()[0].split('_')[-1].lstrip('0')
        if 'Р' in dir and '_' not in dir:
            clean_dir_name = f"Р-{clean_dir_name.replace('Р-', '').lstrip('0')}"


        if clean_dir_name == clean_name(dir_name):
            return os.path.join(base_path, dir)
    return None

for card in cards:
    for el in card['шифры']:
        # Очищаем номера фонда, описи и дела
        el['номер фонда'] = clean_name(el['номер фонда'])
        el['номер описи'] = clean_name(el['номер описи'])
        el['номер дела'] = clean_name(el['номер дела'])
        res_flag = '300'
        # проверяем наличие копий
        if el['разрешение'] == '300':
            print(f'\n\nПроверяем:\n{card["фио"]}\n{el}\n\n')
            SH_path = r'\\192.168.1.185\эл_образы\Единицы хранения\Дореволюционные'
            if el['номер фонда'].startswith("Р"):
                SH_path = r'\\192.168.1.185\эл_образы\Единицы хранения\Советский период'

            # Поиск фонда
            fund_path = find_directory(SH_path, el['номер фонда'])
            if fund_path:
                print(f'Найден фонд {fund_path}')

                # Поиск описи
                op_path = find_directory(fund_path, el['номер описи'])
                if op_path:
                    print(f'Найдена опись {op_path}')

                    # Поиск дела
                    d_path = find_directory(op_path, el['номер дела'])
                    if d_path:
                        print(f'Найдена папка дела: {d_path}')
                        el['разрешение'] = 'копии'
                        res_current = 'копии'
                    else:
                        print(f"Дело {el['номер фонда']}_{el['номер описи']}_{el['номер дела']} не найдено")
                else:
                    print(f"Опись {el['номер фонда']}_{el['номер описи']} не найдена")
            else:
                print(f"Фонд {el['номер фонда']} не найден")

        # -------------------
        # создаём папки
        # -------------------
        base_path = f"{card['номер каиса']} {card['фио']}"
        if el['разрешение'] == 'копии':
            base_path += "_копии"

        # гарантируем, что папка под конкретное разрешение существует
        if not os.path.exists(base_path):
            os.mkdir(base_path)

        # подпапка под шифр
        shifr_path = os.path.join(base_path, f"{el['номер фонда']}_{el['номер описи']}_{el['номер дела']}")
        if not os.path.exists(shifr_path):
            os.mkdir(shifr_path)
            print(f"\n\nСоздана папка:   {shifr_path}\n")

sorted_cards = sorted(cards, key=lambda x: int(x['номер каиса']))

print(f'\nИзмененный список заказов:\n')
for el in sorted_cards:
    print(el)
# #
# # требование
# #


def add_text_to_cell(cell, text, font_name='Times New Roman', font_size=12):
    if cell.paragraphs:
        p = cell.paragraphs[0]
        p.text = text
        for run in p.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    else:
        cell.text = text


def set_font_and_alignment(cell, font_name, font_size, alignment):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        paragraph.alignment = alignment


def save_document(doc, base_path, base_name):
    counter = 0
    while True:
        if counter == 0:
            full_path = os.path.join(base_path, f"{base_name}.docx")
        else:
            full_path = os.path.join(base_path, f"{base_name} ({counter}).docx")
        if not os.path.exists(full_path):
            doc.save(full_path)
            return full_path
        counter += 1


# Загрузка документа
doc_path = "d:\\Заказы\\требование.docx"
doc = Document(doc_path)

# Заполнение таблицы информацией из заказов
table = doc.tables[0]

# Индекс строки в таблице Word, начиная с 1 (предполагая, что первая строка заголовка)
row_idx = 1
row_limit = 13  # Ограничение количества строк в таблице (исключая заголовок)

seen = set()


for order in sorted_cards:
    for shifr in order['шифры']:
        if shifr['разрешение'] == 'копии':
            continue  # Пропускаем заказы с разрешением "копии"
        key = (shifr['номер фонда'], shifr['номер описи'], shifr['номер дела'])
        if key in seen:
            continue  # пропускаем дубликат
        seen.add(key)
        if row_idx > row_limit:
            # Сохраняем текущий документ и создаем новый
            date_ = datetime.today().strftime('%Y %m %d')
            print(date_)
            trebovaniya = 'd:\\Заказы\\_Требования'
            tr_name = f'{date_} Запрос на единицы хранения'

            tr_path = save_document(doc, trebovaniya, tr_name)
            print(f"\n\nТребование сохранено по пути {tr_path}")

            # Создаем новый документ и таблицу
            doc = Document(doc_path)
            table = doc.tables[0]
            row_idx = 1  # Сбрасываем индекс строки для нового документа





        if row_idx >= len(table.rows):
            row = table.add_row()
        else:
            row = table.rows[row_idx]

        add_text_to_cell(row.cells[1], shifr['номер фонда'], 'Times New Roman', 12)
        add_text_to_cell(row.cells[2], shifr['номер описи'], 'Times New Roman', 12)
        add_text_to_cell(row.cells[3], shifr['номер дела'], 'Times New Roman', 12)

        row_idx += 1


# Сохраняем последний документ
date_ = datetime.today().strftime('%Y %m %d')
print(date_)
trebovaniya = 'd:\\Заказы\\_Требования'
tr_name = f'{date_} Запрос на единицы хранения'
tr_path = save_document(doc, trebovaniya, tr_name)
print(f"\n\nТребование сохранено по пути {tr_path}")
#
# #
# # сохранение в excel
# #
#
#
excel_path = 'd:\\Заказы\\Новые заказы каиса.xlsx'

wb = openpyxl.load_workbook(excel_path)
ws = wb.active
fill_color_cop = PatternFill(start_color='FFE699', end_color='FFE699', fill_type='solid')
fill_color_600 = PatternFill(start_color='C6EFCE', end_color='C6EFCE', fill_type='solid')

new_rows = []

for card in sorted_cards:
    # Группируем шифры по разрешению
    rez_groups = {}
    for s in card['шифры']:
        rez = s['разрешение']
        if rez not in rez_groups:
            rez_groups[rez] = []
        rez_groups[rez].append(s)

    # Для каждой группы создаем отдельную строку
    for rez, shifrs in rez_groups.items():
        new_card = card.copy()
        new_card['шифры'] = shifrs
        new_card['rez_for_fill'] = rez
        new_rows.append(new_card)

# Теперь записываем в Excel
for i, card in enumerate(new_rows, start=2):
    shifr_strs = [f"{s['номер фонда']}_{s['номер описи']}_{s['номер дела']}" for s in card['шифры']]
    listy_strs = [s['листы'] for s in card['шифры']]

    ws.cell(row=i, column=3, value="; ".join(shifr_strs))
    cell4 = ws.cell(row=i, column=4)
    cell5 = ws.cell(row=i, column=5)
    cell6 = ws.cell(row=i, column=6)
    ws.cell(row=i, column=7, value=card['фио'])
    ws.cell(row=i, column=8, value=f"каиса {card['номер каиса']}")
    ws.cell(row=i, column=9, value="платно")
    ws.cell(row=i, column=13, value="; ".join(listy_strs))

    # Заливка по разрешению группы
    if card['rez_for_fill'] == 'копии':
        for c in (cell4, cell5, cell6):
            c.fill = fill_color_cop
    elif card['rez_for_fill'] == '600':
        for c in (cell4, cell5, cell6):
            c.fill = fill_color_600

wb.save(excel_path)
